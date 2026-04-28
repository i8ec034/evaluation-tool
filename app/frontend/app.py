import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

import streamlit as st
import pandas as pd
from app.modules.auth_module import authenticate
from app.modules.db_module import SessionLocal, Document, QA, Chunk, QuizResult, QuizAttempt, UserStats
from app.modules.llm_module import generate_qa
from app.modules.vector_db_module import chunk_text, store_embeddings
from app.config.settings import QUESTIONS_PER_QUIZ, THRESHOLD, QA_PER_CHUNK, INTERACTIVE_MODE, CHUNK_SIZE, OVERLAP, CHUNK_TOKEN_THRESHOLD, MAX_ATTEMPTS_PER_USER_PER_DOMAIN
import random
import json
import traceback

st.set_page_config(page_title="Evaluation Tool", layout="wide")

st.markdown(
    "<style>"
    "body, .stApp, .main {background: linear-gradient(180deg, #e8eefb 0%, #ffffff 100%) !important;}"
    ".css-18e3th9 {background: transparent !important;}"
    "[data-testid='stSidebar'] {background: transparent !important;}"
    "</style>", unsafe_allow_html=True)

# ======= GAMIFICATION HELPERS =======
def get_stars(score):
    """Get star rating based on percentage score"""
    if score >= 0.95:
        return "⭐⭐⭐⭐⭐"
    elif score >= 0.85:
        return "⭐⭐⭐⭐"
    elif score >= 0.70:
        return "⭐⭐⭐"
    elif score >= 0.50:
        return "⭐⭐"
    else:
        return "⭐"

def check_badges(username, score, correct_count, total, domain):
    """Check and award badges"""
    db = SessionLocal()
    badges_earned = []
    
    try:
        user_stats = db.query(UserStats).filter(UserStats.username == username).first()
        if not user_stats:
            user_stats = UserStats(username=username)
            db.add(user_stats)
        
        existing_badges = json.loads(user_stats.badges) if user_stats.badges else []
        
        # Badge 1: First Quiz
        if "🎯 First Quiz" not in existing_badges:
            badges_earned.append("🎯 First Quiz")
            existing_badges.append("🎯 First Quiz")
        
        # Badge 2: Perfect Score
        if score >= 1.0 and "💯 Perfect Score" not in existing_badges:
            badges_earned.append("💯 Perfect Score")
            existing_badges.append("💯 Perfect Score")
        
        # Badge 3: Master (90%+)
        if score >= 0.90 and "🏆 Quiz Master" not in existing_badges:
            badges_earned.append("🏆 Quiz Master")
            existing_badges.append("🏆 Quiz Master")
        
        # Badge 4: Speed (all correct)
        if correct_count == total and "⚡ Speed Demon" not in existing_badges:
            badges_earned.append("⚡ Speed Demon")
            existing_badges.append("⚡ Speed Demon")
        
        # Badge 5: Consistent (3 quizzes passed)
        passed_count = db.query(QuizResult).filter(
            QuizResult.username == username,
            QuizResult.passed == 1
        ).count()
        if passed_count >= 3 and "🎪 Consistent Learner" not in existing_badges:
            badges_earned.append("🎪 Consistent Learner")
            existing_badges.append("🎪 Consistent Learner")
        
        # Update user stats
        user_stats.total_quizzes = db.query(QuizResult).filter(QuizResult.username == username).count() + 1
        user_stats.total_correct += correct_count
        user_stats.total_questions += total
        user_stats.highest_score = max(user_stats.highest_score or 0, score)
        user_stats.badges = json.dumps(existing_badges)
        
        db.commit()
        db.close()
        return badges_earned
    except Exception as e:
        db.close()
        return []

def get_leaderboard(limit=10):
    """Get top scorers"""
    db = SessionLocal()
    try:
        results = db.query(
            QuizResult.username,
            (QuizResult.score * 100).label('percentage'),
            QuizResult.area
        ).order_by(QuizResult.score.desc()).limit(limit).all()
        db.close()
        return results
    except:
        db.close()
        return []

st.set_page_config(page_title="Evaluation Tool", layout="wide")

if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.role = None
    st.session_state.username = None

def login():
    col1, col2 = st.columns([3, 2])
    with col1:
        st.markdown(
            "<div style='display: flex; align-items: center; justify-content: center; min-height: 320px;'>"
            "<div style='text-align: center;'>"
            "<div style='font-size: 48px; font-weight: 900; letter-spacing: 0.5px; color: #0f2a5c; line-height: 1.1;'>VM Internal Transition Evaluation</div>"
            "</div>"
            "</div>", unsafe_allow_html=True)
    with col2:
        st.markdown("<div style='text-align: center; margin-bottom: 24px;'><span style='font-size: 34px; font-weight: 800; color: #192a56;'>Login</span></div>", unsafe_allow_html=True)
        role = st.selectbox("Role", ["admin", "user", "guest"], index=1)
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")

        if st.button("Login"):
            if authenticate(username, password, role):
                st.session_state.logged_in = True
                st.session_state.role = role
                st.session_state.username = username
                st.success("Logged in!")
                st.rerun()
            else:
                st.error("Invalid credentials")
def admin_panel():
    st.title("Admin Panel")
    
    try:
        admin_option = st.radio("Select Option", ["Upload Document", "Manage Content", "User Management", "Quiz Results", "User Quiz Reset"], horizontal=True)
        
        if admin_option == "Upload Document":
            # Clear session state if switching to upload
            if st.session_state.get("admin_option") != "Upload Document":
                for key in ["selected_domain", "selected_subdomain", "current_domain"]:
                    if key in st.session_state:
                        del st.session_state[key]
            st.session_state.admin_option = "Upload Document"
            
            try:
                from app.modules.auth_module import load_domains, load_subdomains, add_domain, add_subdomain_to_domain
                
                st.markdown("#### 📋 Select or Create Domain")
                domain_list = load_domains()
                
                col1, col2 = st.columns(2)
                with col1:
                    selected_domain = st.selectbox("Select Domain", domain_list or [], key="domain_select")
                    if selected_domain:
                        st.session_state.selected_domain = selected_domain
                        domain = selected_domain
                with col2:
                    new_domain = st.text_input("Create New Domain", placeholder="Type new domain name", key="new_domain_input")
                
                if new_domain:
                    domain = new_domain.strip()
                    if domain and add_domain(domain):
                        st.success(f"✅ Domain '{domain}' added!")
                        st.session_state.selected_domain = domain
                        # Don't rerun immediately
                        # st.rerun()
                
                if st.session_state.get("selected_domain"):
                    domain = st.session_state.selected_domain
                elif not domain and not new_domain:
                    st.warning("⚠️ Please select or create a domain")
                    return

                # Clear subdomain selection if domain changed
                if st.session_state.get("current_domain") != domain:
                    if "selected_subdomain" in st.session_state:
                        del st.session_state.selected_subdomain
                    st.session_state.current_domain = domain

                st.markdown("#### 🏢 Select or Create Subdomain / Team")
                subdomain_list = load_subdomains(domain)
                col3, col4 = st.columns(2)
                with col3:
                    current_selection = st.session_state.get("selected_subdomain", "")
                    options = [""] + (subdomain_list or [])
                    if current_selection and current_selection not in options:
                        options.append(current_selection)
                    selected_subdomain = st.selectbox("Select Subdomain / Team", options, 
                                                    index=options.index(current_selection) if current_selection in options else 0,
                                                    key="subdomain_select")
                    if selected_subdomain and selected_subdomain != "":
                        st.session_state.selected_subdomain = selected_subdomain
                with col4:
                    new_subdomain = st.text_input("Create New Subdomain / Team", placeholder="Type team name", key="new_subdomain_input")
                
                if new_subdomain:
                    subdomain = new_subdomain.strip()
                    if subdomain and add_subdomain_to_domain(domain, subdomain):
                        st.success(f"✅ Subdomain '{subdomain}' added to '{domain}'!")
                        st.session_state.selected_subdomain = subdomain
                        # Don't rerun immediately, let the user see the success message
                        # st.rerun()
                
                # Get subdomain from session state
                subdomain = st.session_state.get("selected_subdomain")
                
                if subdomain:
                    st.info(f"📍 Selected: **Domain:** {domain} | **Subdomain:** {subdomain}")
                else:
                    st.warning("⚠️ Please select or create a subdomain/team")
                    return

                st.markdown("#### 📄 Upload and Process Document")
                uploaded_file = st.file_uploader("Upload Document", type=["pdf", "txt", "docx"])
                if uploaded_file and st.button("Process Document"):
                    with st.spinner("📖 Processing document..."):
                        try:
                            if uploaded_file.type == "text/plain":
                                content = uploaded_file.read().decode("utf-8")
                            elif uploaded_file.type == "application/pdf":
                                from PyPDF2 import PdfReader
                                reader = PdfReader(uploaded_file)
                                content = "".join([page.extract_text() for page in reader.pages])
                            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                from docx import Document as DocxDocument
                                docx_doc = DocxDocument(uploaded_file)
                                content = "\n".join([para.text for para in docx_doc.paragraphs])
                            else:
                                st.error("❌ Unsupported file type")
                                return
                            
                            if not content or len(content.strip()) == 0:
                                st.error("❌ Document is empty")
                                return
                            
                            source_name = getattr(uploaded_file, "name", "uploaded")
                            chunks = chunk_text(content)
                            structured_qa = generate_qa(content)
                            
                            if not structured_qa:
                                st.error("❌ Failed to generate Q&A. Check your API key.")
                                return

                            st.session_state.pending_document = {
                                "domain": domain,
                                "subdomain": subdomain,
                                "source": source_name,
                                "content": content,
                                "chunks": chunks,
                                "structured_qa": structured_qa
                            }
                            st.success("✅ Document processed. Review generated Q&A below before saving.")
                        except Exception as e:
                            st.error(f"❌ Processing error: {str(e)}")
                            st.write(traceback.format_exc())

                if st.session_state.get("pending_document"):
                    pending = st.session_state.pending_document
                    st.markdown("### 📝 Review generated Q&A")
                    st.info("Review each chunk and select or edit only the QA you want to save.")
                    st.write(f"**Domain:** {pending['domain']} | **Subdomain/Team:** {pending['subdomain']}")
                    st.write(f"**File:** {pending['source']} | **Chunks:** {len(pending['chunks'])}")
                    total_generated = sum(len(chunk.get('qas', [])) for chunk in pending['structured_qa'])
                    st.write(f"**Generated QA:** {total_generated}")
                    st.markdown("---")

                    for chunk in pending['structured_qa']:
                        with st.expander(f"Chunk {chunk.get('chunk_index', 0)} | Generated QA: {len(chunk.get('qas', []))}"):
                            preview = chunk.get('chunk_text', '')[:250]
                            st.write(f"**Chunk preview:** {preview}{'...' if len(chunk.get('chunk_text', '')) > 250 else ''}")
                            if not chunk.get('qas'):
                                st.warning("No Q&A generated for this chunk.")
                            for q_idx, qa in enumerate(chunk.get('qas', [])):
                                qa_key = f"pending_qa_{chunk.get('chunk_index', 0)}_{q_idx}"
                                keep_key = qa_key + "_keep"
                                question_key = qa_key + "_question"
                                options_key = qa_key + "_options"
                                answer_key = qa_key + "_answer"
                                qa_options = qa.get('options') or []
                                if not isinstance(qa_options, list):
                                    qa_options = []
                                with st.container():
                                    st.checkbox("Keep QA", value=True, key=keep_key)
                                    st.text_input("Question", value=qa.get('question', ''), key=question_key)
                                    st.text_area("Options (comma-separated)", value=", ".join(qa_options), key=options_key)
                                    st.text_input("Answer", value=qa.get('answer', ''), key=answer_key)
                                    st.markdown("---")

                    if st.button("Save approved Q&A"):
                        approved_items = []
                        for chunk in pending['structured_qa']:
                            for q_idx, qa in enumerate(chunk.get('qas', [])):
                                qa_key = f"pending_qa_{chunk.get('chunk_index', 0)}_{q_idx}"
                                if not st.session_state.get(qa_key + "_keep", True):
                                    continue
                                question = st.session_state.get(qa_key + "_question", "").strip()
                                answer = st.session_state.get(qa_key + "_answer", "").strip()
                                options_text = st.session_state.get(qa_key + "_options", "")
                                options = [opt.strip() for opt in options_text.split(",") if opt.strip()]
                                if not question or not answer or not options:
                                    continue
                                approved_items.append({
                                    "chunk_index": chunk.get('chunk_index', 0),
                                    "question": question,
                                    "options": options,
                                    "answer": answer
                                })

                        try:
                            db = SessionLocal()
                            doc = Document(
                                area=pending['domain'],
                                domain=pending['domain'],
                                subdomain=pending['subdomain'],
                                content=pending['content'],
                                source=pending['source']
                            )
                            db.add(doc)
                            db.commit()
                            db.refresh(doc)

                            # Try to store embeddings (optional - will continue if it fails)
                            try:
                                store_embeddings(
                                    pending['domain'],
                                    pending['chunks'],
                                    document_id=doc.id,
                                    source=pending['source'],
                                    domain=pending['domain'],
                                    subdomain=pending['subdomain']
                                )
                            except Exception as embed_error:
                                st.warning(f"⚠️ Embeddings storage failed: {str(embed_error)}")
                                st.info("✅ QA saved successfully - embeddings are optional for quiz functionality")

                            total_saved = 0
                            for i, chunk_content in enumerate(pending['chunks']):
                                chunk_obj = Chunk(
                                    document_id=doc.id,
                                    index=i,
                                    content=chunk_content,
                                    source=pending['source'],
                                    metadata_json=json.dumps({
                                        "domain": pending['domain'],
                                        "subdomain": pending['subdomain'],
                                        "area": pending['domain'],
                                        "chunk_index": i,
                                        "source": pending['source']
                                    })
                                )
                                db.add(chunk_obj)
                                db.flush()

                                for qa_item in approved_items:
                                    if qa_item['chunk_index'] != i:
                                        continue
                                    qa_obj = QA(
                                        document_id=doc.id,
                                        chunk_id=chunk_obj.id,
                                        question=qa_item['question'],
                                        options=json.dumps(qa_item['options']),
                                        answer=qa_item['answer']
                                    )
                                    db.add(qa_obj)
                                    total_saved += 1

                            db.commit()
                            db.close()
                            del st.session_state.pending_document
                            st.success(f"✅ Saved document and {total_saved} approved Q&A items.")
                        except Exception as e:
                            st.error(f"❌ Save error: {str(e)}")
                            st.write(traceback.format_exc())
            except Exception as e:
                st.error(f"❌ Upload error: {str(e)}")
        
        elif admin_option == "Manage Content":
            try:
                st.markdown("### 🔎 Domain Content Overview")
                db = SessionLocal()
                
                domains = db.query(Document.domain).filter(Document.domain != None).distinct().all()
                domain_list = [d[0] for d in domains if d[0]]
                
                if not domain_list:
                    st.info("📭 No domains with content yet")
                    db.close()
                    return
                
                selected_domain = st.selectbox("Select Domain", domain_list)
                st.markdown(f"#### 📊 Content for: **{selected_domain}**")
                st.markdown("### 🧾 Domain Summary")
                st.write(f"- **Chunk split threshold:** {CHUNK_TOKEN_THRESHOLD} tokens")
                st.write(f"- **Chunk size:** {CHUNK_SIZE} characters")
                st.write(f"- **Chunk overlap:** {OVERLAP} characters")
                st.write("- **Metadata storage:** Each chunk stores `domain`, `subdomain`, `area`, `chunk_index`, and `source` in `chunks.metadata_json`; embeddings also include these fields in ChromaDB.")
                st.markdown("---")
                
                subdomain_list = [s[0] for s in db.query(Document.subdomain).filter(Document.domain == selected_domain, Document.subdomain != None).distinct().all() if s[0]]
                selected_subdomain = st.selectbox("Select Subdomain / Team", ["All"] + subdomain_list) if subdomain_list else "All"
                
                query = db.query(Document).filter(Document.domain == selected_domain)
                if selected_subdomain and selected_subdomain != "All":
                    query = query.filter(Document.subdomain == selected_subdomain)
                docs = query.all()
                total_qas = db.query(QA).join(Document).filter(Document.domain == selected_domain)
                if selected_subdomain and selected_subdomain != "All":
                    total_qas = total_qas.filter(Document.subdomain == selected_subdomain)
                total_qas = total_qas.count()
                
                st.markdown(f"- **Documents:** {len(docs)} | **Q&A:** {total_qas}")
                st.markdown("---")
                
                with st.expander("📄 Documents"):
                    if docs:
                        for doc in docs:
                            chunk_count = len(doc.chunks)
                            qas_count = db.query(QA).filter(QA.document_id == doc.id).count()
                            chunked_text = "Yes" if chunk_count > 1 else "No"
                            with st.expander(f"Doc {doc.id} | Chunks: {chunk_count} | Q&A: {qas_count} | Chunked: {chunked_text}"):
                                st.write(f"**Source Document:** {doc.source or 'uploaded'}")
                                st.write(f"**Domain:** {doc.domain}")
                                st.write(f"**Chunked:** {chunked_text}")
                                if doc.chunks:
                                    for chunk in doc.chunks:
                                        chunk_qas_count = db.query(QA).filter(QA.chunk_id == chunk.id).count()
                                        try:
                                            chunk_metadata = json.loads(chunk.metadata_json) if chunk.metadata_json else {}
                                        except Exception:
                                            chunk_metadata = {"raw": chunk.metadata_json}
                                        with st.expander(f"Chunk {chunk.index} | Q&A: {chunk_qas_count} | Source: {chunk.source or 'unknown'}"):
                                            st.write(f"**Metadata:** {chunk_metadata}")
                                            preview = chunk.content[:250]
                                            st.write(f"**Text Preview:** {preview}{'...' if len(chunk.content) > 250 else ''}")
                                else:
                                    st.info("No chunk data available")
                                col1, col2 = st.columns([4, 1])
                                with col2:
                                    if st.button("Delete", key=f"del_doc_{doc.id}"):
                                        db.query(QA).filter(QA.document_id == doc.id).delete()
                                        db.query(Chunk).filter(Chunk.document_id == doc.id).delete()
                                        db.delete(doc)
                                        db.commit()
                                        st.success("✅ Deleted")
                                        st.rerun()
                    else:
                        st.info("No documents")
                
                with st.expander("❓ Q&A Preview"):
                    if total_qas > 0:
                        displayed = 0
                        for doc in docs:
                            if displayed >= 10:
                                break
                            qas = db.query(QA).filter(QA.document_id == doc.id).limit(10 - displayed).all()
                            for qa in qas:
                                displayed += 1
                                try:
                                    options = json.loads(qa.options)
                                except Exception:
                                    options = []
                                if not isinstance(options, list):
                                    options = []
                                st.write(f"{displayed}. {qa.question}")
                                st.write(f"Answer: {qa.answer}")
                                st.write("---")
                                if displayed >= 10:
                                    break
                    else:
                        st.info("No Q&A")
                
                db.close()
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
        
        elif admin_option == "User Management":
            try:
                from app.modules.auth_module import get_all_users, add_admin_user, add_user, add_guest_user, remove_user, load_domains, load_subdomains
                from app.config.settings import USE_DB_FOR_AUTH
                
                if not USE_DB_FOR_AUTH:
                    st.warning("⚠️ User management is only available when USE_DB_FOR_AUTH is set to True in settings.py")
                    return
                
                st.markdown("### 👥 User Management")
                
                # Display current users
                st.markdown("#### 📋 Current Users")
                users = get_all_users()
                if users:
                    for user in users:
                        col1, col2, col3, col4 = st.columns([3, 2, 3, 1])
                        with col1:
                            st.write(f"**{user['username']}**")
                        with col2:
                            st.write(f"Role: {user['role']}")
                        with col3:
                            if user['role'] == 'user':
                                st.write(f"Domain: {user.get('domain', 'N/A')}")
                                st.write(f"Subdomain: {user.get('subdomain', 'N/A')}")
                            else:
                                st.write("-")
                        with col4:
                            if st.button("🗑️", key=f"del_user_{user['username']}_{user['role']}"):
                                if remove_user(user['username'], user['role']):
                                    st.success(f"✅ Removed {user['username']}")
                                    st.rerun()
                                else:
                                    st.error("❌ Failed to remove user")
                else:
                    st.info("No users found")
                
                st.markdown("---")
                st.markdown("#### ➕ Add New User")
                
                user_type = st.selectbox("User Type", ["admin", "user", "guest"], key="user_type")
                
                col1, col2 = st.columns(2)
                with col1:
                    new_username = st.text_input("Username", key="new_username")
                with col2:
                    new_password = st.text_input("Password", type="password", key="new_password")
                
                if user_type == "user":
                    domains = load_domains()
                    selected_domain = st.selectbox("Domain", domains or [], key="user_domain")
                    if selected_domain:
                        subdomains = load_subdomains(selected_domain)
                        selected_subdomain = st.selectbox("Subdomain", subdomains or [], key="user_subdomain")
                    else:
                        selected_subdomain = None
                        st.warning("Please select a domain first")
                
                if st.button("Add User"):
                    if not new_username or not new_password:
                        st.error("❌ Username and password are required")
                        return
                    
                    success = False
                    if user_type == "admin":
                        success = add_admin_user(new_username, new_password)
                    elif user_type == "user":
                        if not selected_domain or not selected_subdomain:
                            st.error("❌ Domain and subdomain are required for users")
                            return
                        success = add_user(new_username, new_password, selected_domain, selected_subdomain)
                    elif user_type == "guest":
                        success = add_guest_user(new_username, new_password)
                    
                    if success:
                        st.success(f"✅ Added {user_type} user: {new_username}")
                        st.rerun()
                    else:
                        st.error(f"❌ Failed to add {user_type} user")
                        
            except Exception as e:
                st.error(f"❌ User management error: {str(e)}")
        
        elif admin_option == "Quiz Results":
            try:
                st.markdown("### 📊 Quiz Results")
                db = SessionLocal()

                results = db.query(QuizResult).order_by(QuizResult.score.desc()).all()
                if not results:
                    st.info("📭 No quiz results yet")
                    db.close()
                else:
                    st.markdown(f"**Total Results:** {len(results)}")
                    for res in results:
                        col1, col2, col3 = st.columns([4, 2, 1])
                        with col1:
                            st.write(f"**{res.username}** | {res.area} | {res.score*100:.1f}% | {'✅' if res.passed else '❌'}")
                        with col2:
                            st.write(f"Total: {res.total} questions")
                        with col3:
                            if st.button("Delete", key=f"del_result_{res.id}"):
                                db.delete(res)
                                db.commit()
                                st.success("✅ Deleted")
                                st.rerun()
                    db.close()
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
        elif admin_option == "User Quiz Reset":
            try:
                st.markdown("### 🔄 Reset User Attempts")
                db = SessionLocal()
                
                attempts = db.query(QuizAttempt.domain).distinct().all()
                domain_list = [a[0] for a in attempts]
                
                if not domain_list:
                    st.info("📭 No attempts recorded")
                    db.close()
                else:
                    selected_domain = st.selectbox("Select Domain", domain_list, key="reset_domain")
                    domain_attempts = db.query(QuizAttempt).filter(QuizAttempt.domain == selected_domain).all()
                    
                    st.write(f"**Total Attempts:** {len(domain_attempts)}")
                    
                    if domain_attempts:
                        for attempt in domain_attempts:
                            col1, col2 = st.columns([4, 1])
                            with col1:
                                st.write(f"👤 **{attempt.username}**")
                            with col2:
                                if st.button("Reset", key=f"reset_{attempt.id}"):
                                    db.delete(attempt)
                                    db.commit()
                                    st.success("✅ Reset")
                                    st.rerun()
                    db.close()
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
    except Exception as e:
        st.error(f"❌ Admin error: {str(e)}")

def user_quiz():
    try:
        db = SessionLocal()
        df_user = pd.read_excel('data/users.xlsx', sheet_name='users')
        user_row = df_user[df_user['username'] == st.session_state.username]
        
        if user_row.empty:
            st.error("❌ User not found")
            return
        
        domain = user_row['domain'].iloc[0]
        subdomain = user_row['subdomain'].iloc[0] if 'subdomain' in user_row.columns else None
        st.title(f" Evaluation - {domain}")
        st.markdown("---")

        # Check attempt count for this domain/subdomain
        attempt_query = db.query(QuizAttempt).filter(
            QuizAttempt.username == st.session_state.username,
            QuizAttempt.domain == domain
        )
        if subdomain and str(subdomain).strip():
            attempt_query = attempt_query.filter(QuizAttempt.subdomain == subdomain)

        attempt_count = attempt_query.count()
        if attempt_count >= MAX_ATTEMPTS_PER_USER_PER_DOMAIN:
            st.warning(f"⚠️ Maximum attempts reached ({MAX_ATTEMPTS_PER_USER_PER_DOMAIN}).")
            st.info("Contact admin to reset your attempts.")
            db.close()
            return

        remaining = MAX_ATTEMPTS_PER_USER_PER_DOMAIN - attempt_count
        st.info(f"You have {remaining} attempt{'s' if remaining != 1 else ''} remaining for domain '{domain}'")

        # Get Q&A for the user's domain and subdomain
        docs = db.query(Document).filter(Document.domain == domain)
        if subdomain and str(subdomain).strip():
            docs = docs.filter(Document.subdomain == subdomain)
        docs = docs.all()
        qas = []
        for doc in docs:
            qas.extend(db.query(QA).filter(QA.document_id == doc.id).all())
        
        if not qas:
            st.warning(f"📚 No content for {domain} yet")
            db.close()
            return
        
        st.markdown("### Begin Challenge")
        if st.button("Start Quiz"):
            selected = random.sample(qas, min(QUESTIONS_PER_QUIZ, len(qas)))
            st.session_state.quiz = selected
            st.session_state.current_q = 0
            st.session_state.score = 0
            st.session_state.streak = 0
            st.session_state.answer_feedback = None
            st.session_state.quiz_started = True
            st.rerun()

        if st.session_state.get('quiz_started', False) and 'quiz' in st.session_state:
            if st.session_state.current_q < len(st.session_state.quiz):
                q = st.session_state.quiz[st.session_state.current_q]
                st.progress(st.session_state.current_q / len(st.session_state.quiz))
                
                # Show combo streak
                st.markdown(f"### Q {st.session_state.current_q + 1}/{len(st.session_state.quiz)} | 🔥 Streak: {st.session_state.get('streak', 0)}")
                st.markdown(f"#### {q.question}")
                
                try:
                    options = json.loads(q.options)
                except Exception:
                    options = []
                if not isinstance(options, list):
                    options = []
                if not options:
                    options = ["No options available"]
                
                # Show feedback if already answered
                if st.session_state.get('answer_feedback') and st.session_state.answer_feedback['question_index'] == st.session_state.current_q:
                    feedback = st.session_state.answer_feedback
                    if feedback['is_correct']:
                        st.success(f"✅ Correct! Your answer: **{feedback['user_answer']}** | 🔥 Streak: {st.session_state.streak}")
                    else:
                        st.error(f"❌ Wrong! Your answer: **{feedback['user_answer']}** | Correct answer: **{feedback['correct_answer']}**")
                    
                    if st.button("⏭️ Next Question"):
                        st.session_state.answer_feedback = None
                        st.session_state.current_q += 1
                        st.rerun()
                else:
                    # Show answer options
                    answer = st.radio("Select your answer:", options, key=f"q{st.session_state.current_q}")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        if st.button("✅ Submit Answer"):
                            is_correct = answer == q.answer
                            if is_correct:
                                st.session_state.score += 1
                                st.session_state.streak = st.session_state.get('streak', 0) + 1
                            else:
                                st.session_state.streak = 0  # Reset streak
                            
                            st.session_state.answer_feedback = {
                                'question_index': st.session_state.current_q,
                                'user_answer': answer,
                                'correct_answer': q.answer,
                                'is_correct': is_correct
                            }
                            st.rerun()
                    with col2:
                        if st.button("⏭️ Skip"):
                            st.session_state.current_q += 1
                            st.rerun()
            else:
                score = st.session_state.score
                total = len(st.session_state.quiz)
                percentage = score / total
                passed = 1 if percentage >= THRESHOLD else 0
                
                st.markdown("---")
                st.markdown("### 🏆 Complete!")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Score", f"{score}/{total}")
                with col2:
                    st.metric("%", f"{percentage*100:.1f}%")
                with col3:
                    st.markdown(f"<div style='text-align: center;'><h3>{get_stars(percentage)}</h3></div>", unsafe_allow_html=True)
                
                st.markdown("---")
                if passed:
                    st.success(f"🎉 You passed with {percentage*100:.1f}%!")
                    st.balloons()
                else:
                    st.error(f"Need {THRESHOLD*100:.0f}%. You got {percentage*100:.1f}%")
                
                # Check and award badges
                badges = check_badges(st.session_state.username, percentage, score, total, domain)
                if badges:
                    st.markdown("### 🎊 New Badges Earned!")
                    for badge in badges:
                        st.write(f"✨ {badge}")
                
                try:
                    db = SessionLocal()
                    result = QuizResult(
                        username=st.session_state.username,
                        area=domain,
                        subdomain=subdomain,
                        score=percentage,
                        total=total,
                        passed=passed
                    )
                    db.add(result)
                    db.commit()
                    
                    attempt = QuizAttempt(
                        username=st.session_state.username,
                        domain=domain,
                        subdomain=subdomain
                    )
                    db.add(attempt)
                    db.commit()
                    st.success("✅ Saved!")
                except Exception as e:
                    st.error(f"❌ Save error: {str(e)}")
                finally:
                    db.close()
    except Exception as e:
        st.error(f"❌ Quiz error: {str(e)}")

def guest_view():
    try:
        st.title("👤 Guest View")
        db = SessionLocal()
        docs = db.query(Document).all()
        
        if not docs:
            st.info("📚 No content")
            db.close()
            return
        
        areas = list(set([doc.area for doc in docs]))
        area = st.selectbox("Select Area", areas)
        
        docs = db.query(Document).filter(Document.area == area).all()
        qas = []
        for doc in docs:
            qas.extend(db.query(QA).filter(QA.document_id == doc.id).all())
        
        st.markdown(f"### 📖 {area}")
        for i, qa in enumerate(qas[:20], 1):
            st.write(f"{i}. {qa.question}")
            try:
                options = json.loads(qa.options)
            except Exception:
                options = []
            if not isinstance(options, list):
                options = []
            st.write(f"Answer: {qa.answer}")
            st.write("---")
        
        db.close()
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")

def show_leaderboard():
    """Display top 10 scorers leaderboard"""
    try:
        st.title("🏆 Top 10 Scorers")
        results = get_leaderboard(10)
        
        if not results:
            st.info("📭 No scores yet")
            return
        
        # Display as table
        data = []
        for idx, (username, percentage, area) in enumerate(results, 1):
            stars = get_stars(percentage / 100)
            data.append({
                "Rank": f"#{idx}",
                "Player": username,
                "Score": f"{percentage:.1f}%",
                "Rating": stars,
                "Domain": area
            })
        
        df = pd.DataFrame(data)
        st.dataframe(df, use_container_width=True, hide_index=True)
        
        # Show top player stats
        if results:
            top_user = results[0][0]
            db = SessionLocal()
            user_stats = db.query(UserStats).filter(UserStats.username == top_user).first()
            
            st.markdown("---")
            st.markdown(f"### 👑 Leader: **{top_user}**")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Quizzes Taken", user_stats.total_quizzes if user_stats else 0)
            with col2:
                highest_score = (user_stats.highest_score if user_stats and user_stats.highest_score is not None else 0)
                st.metric("Highest Score", f"{highest_score*100:.1f}%")
            with col3:
                badge_count = 0
                if user_stats and user_stats.badges:
                    badges = json.loads(user_stats.badges)
                    badge_count = len(badges)
                st.metric("Badges", badge_count)
            
            # Show badges
            if user_stats and user_stats.badges:
                badges = json.loads(user_stats.badges)
                st.markdown("**Badges Earned:**")
                badge_text = " ".join(badges)
                st.write(badge_text)
            
            db.close()
    except Exception as e:
        st.error(f"❌ Leaderboard error: {str(e)}")

def show_user_profile():
    """Show current user's profile with badges"""
    try:
        db = SessionLocal()
        user_stats = db.query(UserStats).filter(UserStats.username == st.session_state.username).first()
        
        st.sidebar.markdown("---")
        st.sidebar.markdown("### 📊 Your Stats")
        
        if user_stats:
            col1, col2 = st.sidebar.columns(2)
            with col1:
                st.metric("Quizzes", user_stats.total_quizzes)
            with col2:
                st.metric("Best Score", f"{(user_stats.highest_score or 0)*100:.1f}%")
            
            if user_stats.badges:
                badges = json.loads(user_stats.badges)
                st.sidebar.markdown("**Badges:**")
                for badge in badges:
                    st.sidebar.write(badge)
        
        db.close()
    except Exception as e:
        pass

if not st.session_state.logged_in:
    login()
else:
    st.sidebar.markdown("---")
    if st.sidebar.button("🚪 Logout"):
        st.session_state.logged_in = False
        st.rerun()
    
    st.sidebar.write(f"**{st.session_state.username}** ({st.session_state.role})")
    
    # Show leaderboard option for all users
    show_leaderboard_option = st.sidebar.checkbox("📊 View Leaderboard")
    
    if show_leaderboard_option:
        show_leaderboard()
    elif st.session_state.role == "admin":
        admin_panel()
    elif st.session_state.role == "user":
        show_user_profile()
        user_quiz()
    else:
        guest_view()
